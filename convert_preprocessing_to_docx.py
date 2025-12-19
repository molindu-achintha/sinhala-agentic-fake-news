#!/usr/bin/env python3
"""
Convert preprocessing pipeline explanation to Word document
"""
from docx import Document
from docx.shared import Pt
import re


def convert_preprocessing_to_docx():
    """Convert preprocessing explanation to Word doc."""
    
    with open('preprocessing_pipeline_explanation.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        if not line.strip():
            i += 1
            continue
        
        # Main headers (=== lines)
        if '=' * 30 in line:
            i += 1
            if i < len(lines) and lines[i].strip() and '=' not in lines[i]:
                heading = lines[i].strip()
                if 'CHAPTER' in heading or len(heading) > 50:
                    doc.add_heading(heading, level=0)
                else:
                    doc.add_heading(heading, level=1)
            i += 1
            while i < len(lines) and '=' * 30 in lines[i]:
                i += 1
            continue
        
        # Section headers (X.X TITLE format)
        section_match = re.match(r'^(\d+\.\d*)\s+(.+)$', line.strip())
        if section_match:
            doc.add_heading(line.strip(), level=2)
            i += 1
            continue
        
        # Subsection headers (ending with ---)
        if i + 1 < len(lines) and lines[i + 1].strip().startswith('---'):
            doc.add_heading(line.strip(), level=3)
            i += 2
            continue
        
        # Skip standalone --- lines
        if line.strip().startswith('---') or line.strip().endswith('---'):
            i += 1
            continue
        
        # Code blocks (```)
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue
        
        # Tables (+ or |)
        if line.strip().startswith('+') or line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and (lines[i].strip().startswith('+') or lines[i].strip().startswith('|')):
                table_lines.append(lines[i])
                i += 1
            table_text = '\n'.join(table_lines)
            p = doc.add_paragraph()
            run = p.add_run(table_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue
        
        # Regular paragraph
        if line.strip():
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i]
                if '=' * 20 in next_line or next_line.strip().endswith('---'):
                    break
                if re.match(r'^(\d+\.\d*)\s+', next_line.strip()):
                    break
                if next_line.strip().startswith('```'):
                    break
                if next_line.strip().startswith('+') or next_line.strip().startswith('|'):
                    break
                if not next_line.strip():
                    break
                para_lines.append(next_line)
                i += 1
            
            para_text = '\n'.join(para_lines)
            doc.add_paragraph(para_text)
            continue
        
        i += 1
    
    output_file = "Preprocessing_Pipeline_Thesis.docx"
    doc.save(output_file)
    print(f"Created: {output_file}")


if __name__ == "__main__":
    convert_preprocessing_to_docx()
