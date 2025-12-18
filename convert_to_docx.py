#!/usr/bin/env python3
"""
Convert report.txt to Word document (.docx)
Uses python-docx package
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


def convert_txt_to_docx(txt_path: str, docx_path: str):
    """Convert plain text report to formatted Word document."""
    
    # Read the text file
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines at the start
        if not line.strip():
            i += 1
            continue
        
        # Detect chapter/section headers (lines with ===)
        if '=' * 20 in line:
            # Next line is likely a heading
            i += 1
            if i < len(lines):
                heading_text = lines[i].strip()
                if heading_text and '=' not in heading_text:
                    # Determine heading level
                    if 'CHAPTER' in heading_text.upper():
                        doc.add_heading(heading_text, level=1)
                    elif 'APPENDIX' in heading_text.upper():
                        doc.add_heading(heading_text, level=1)
                    elif heading_text.isupper() or heading_text.startswith('MSc'):
                        doc.add_heading(heading_text, level=1)
                    else:
                        doc.add_heading(heading_text, level=2)
                i += 1
            # Skip closing === line
            while i < len(lines) and '=' * 20 in lines[i]:
                i += 1
            continue
        
        # Detect subsection headers (lines ending with dashes like ---)
        if line.strip().endswith('---') or (line.strip().startswith('-') and len(line.strip()) > 10 and line.strip().count('-') > 5):
            i += 1
            continue
        
        # Detect numbered sections like "1.1 INTRODUCTION"
        section_match = re.match(r'^(\d+\.[\d.]*)\s+([A-Z].*)', line.strip())
        if section_match:
            doc.add_heading(line.strip(), level=2)
            i += 1
            continue
        
        # Detect lettered sections like "I.1 WHAT IS..."
        letter_section_match = re.match(r'^([A-Z]+\.[\d.]+)\s+(.+)', line.strip())
        if letter_section_match:
            doc.add_heading(line.strip(), level=2)
            i += 1
            continue
        
        # Regular paragraph text
        if line.strip():
            # Collect consecutive non-empty lines as one paragraph
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and '=' * 20 not in lines[i] and not lines[i].strip().endswith('---'):
                # Check if this looks like a new section
                if re.match(r'^(\d+\.[\d.]*)\s+[A-Z]', lines[i].strip()):
                    break
                if re.match(r'^([A-Z]+\.[\d.]+)\s+', lines[i].strip()):
                    break
                if lines[i].strip().startswith('CHAPTER') or lines[i].strip().startswith('APPENDIX'):
                    break
                para_lines.append(lines[i])
                i += 1
            
            # Join and add as paragraph
            para_text = '\n'.join(para_lines)
            p = doc.add_paragraph(para_text)
            continue
        
        i += 1
    
    # Save the document
    doc.save(docx_path)
    print(f"Successfully converted to: {docx_path}")


if __name__ == "__main__":
    txt_file = "report.txt"
    docx_file = "MSc_Project_Report_Sinhala_Fake_News_Detection.docx"
    convert_txt_to_docx(txt_file, docx_file)
