#!/usr/bin/env python3
"""
Convert methodology section to Word document
"""
from docx import Document
from docx.shared import Pt
import re


def convert_methodology_to_docx():
    """Convert methodology explanation to Word doc."""
    
    with open('methodology_section.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Add Title
    title = doc.add_heading('CHAPTER 4: METHODOLOGY', level=0)
    title.alignment = 1  # Center
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines at start
        if not line.strip():
            i += 1
            continue
            
        # Skip the manual header lines I wrote in the txt file
        if 'CHAPTER 4:' in line or '=' * 10 in line and i < 10:
            i += 1
            continue
        
        # Main headers (=== lines)
        if '=' * 30 in line:
            # Check the line BEFORE the === line
            if i > 0 and lines[i-1].strip():
                # This case shouldn't happen with my current parser logic
                pass
            i += 1
            continue
            
        # 4.X Headers
        header_match = re.match(r'^(4\.\d+)\s+(.+)$', line.strip())
        if header_match:
            doc.add_heading(line.strip(), level=1)
            i += 1
            # Skip underline if present
            if i < len(lines) and '=' * 10 in lines[i]:
                i += 1
            continue
            
        # 4.X.X Subheaders
        subheader_match = re.match(r'^(4\.\d+\.\d+)\s+(.+)$', line.strip())
        if subheader_match:
            doc.add_heading(line.strip(), level=2)
            i += 1
            # Skip underline if present
            if i < len(lines) and '-' * 10 in lines[i]:
                i += 1
            continue

        # Code blocks (```)
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue
        
        # Blockquotes (>)
        if line.strip().startswith('>'):
            quote_text = line.strip().replace('>', '').strip()
            p = doc.add_paragraph(quote_text)
            p.style = 'Quote' 
            i += 1
            continue
            
        # Regular paragraph
        if line.strip():
            # Check if it's a list item
            if line.strip().startswith('- '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
                i += 1
                continue
                
            # Collect paragraph lines
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i]
                if not next_line.strip():
                    break
                if next_line.strip().startswith('```'):
                    break
                if re.match(r'^4\.\d+', next_line.strip()):
                    break
                if next_line.strip().startswith('- '):
                    break
                    
                para_lines.append(next_line)
                i += 1
            
            para_text = '\n'.join(para_lines)
            doc.add_paragraph(para_text)
            continue
            
        i += 1
    
    output_file = "Methodology_Section_Thesis.docx"
    doc.save(output_file)
    print(f"Created: {output_file}")


if __name__ == "__main__":
    convert_methodology_to_docx()
