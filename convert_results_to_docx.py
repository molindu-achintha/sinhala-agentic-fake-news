#!/usr/bin/env python3
"""
Convert Results section to Word document.
Supports tables.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_results_to_docx():
    """Convert results section text to formatted Word doc."""
    
    with open('results_analysis_section.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines at start
        if not line.strip():
            i += 1
            continue
            
        # Title Block
        if 'CHAPTER 5:' in line:
            p = doc.add_heading(line.strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            while i < len(lines) and (not lines[i].strip() or '=' in lines[i] or 'Sinhala' in lines[i] or 'December' in lines[i]):
                if lines[i].strip() and '=' not in lines[i]:
                     p = doc.add_paragraph(lines[i].strip())
                     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
            continue
        
        # 5.X Headers (===)
        if i + 1 < len(lines) and lines[i+1].strip().startswith('==='):
            doc.add_heading(line.strip(), level=1)
            i += 2
            continue
            
        # 5.X.X Subheaders (---)
        if i + 1 < len(lines) and lines[i+1].strip().startswith('---'):
            doc.add_heading(line.strip(), level=2)
            i += 2
            continue

        # Detect Tables
        # Assuming table starts with "| Model" or similar pipe-based syntax
        if line.strip().startswith('|'):
            # Collect table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Process table
            if len(table_lines) > 2: # Header, Separator, Data
                # Parse header
                headers = [h.strip() for h in table_lines[0].split('|') if h.strip()]
                # content lines (skip [1] which is separator)
                data_rows = table_lines[2:]
                
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for idx, h in enumerate(headers):
                    hdr_cells[idx].text = h
                
                for row_line in data_rows:
                    cols = [c.strip() for c in row_line.split('|') if c.strip() or c == ""] 
                    # Note: split logic might act up if empty cells exist, simplifying for now
                    # Better split: slice[1:-1]
                    raw_cols = row_line.split('|')
                    # remove first and last empty strings from split
                    if raw_cols[0] == '': raw_cols.pop(0)
                    if raw_cols[-1] == '': raw_cols.pop(-1)
                    
                    row_cells = table.add_row().cells
                    for idx, c in enumerate(raw_cols):
                        if idx < len(row_cells):
                            row_cells[idx].text = c.strip().replace('**', '') # remove markdown bold
            continue
        
        # Detect bold table/figure captions
        if line.strip().startswith('**Table') or line.strip().startswith('**Figure'):
             doc.add_paragraph(line.strip().replace('**', ''), style='Caption')
             i += 1
             continue

        # Regular list handling
        if line.strip().startswith('- '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s+', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Blockquote (Explanation samples)
        if line.strip().startswith('>'):
             p = doc.add_paragraph(line.strip().replace('>', ''))
             p.style = 'Quote'
             i += 1
             continue

        # Regular Paragraph
        if line.strip():
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i]
                if not next_line.strip(): break
                if next_line.startswith('='): break
                if next_line.startswith('-'): break
                if next_line.startswith('|'): break
                if next_line.startswith('**'): break
                if re.match(r'^\d+\.\s+', next_line): break
                
                para_lines.append(next_line.strip())
                i += 1
            doc.add_paragraph(' '.join(para_lines))
            continue
            
        i += 1
    
    output_file = "Results_Analysis_Thesis.docx"
    doc.save(output_file)
    print(f"Created: {output_file}")

if __name__ == "__main__":
    convert_results_to_docx()
