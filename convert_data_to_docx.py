#!/usr/bin/env python3
"""
Convert data_section.txt to Word document (.docx)
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def convert_data_section_to_docx():
    """Convert data section text to formatted Word document."""
    
    # Read the text file
    with open('data_section.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Main chapter headers (lines with ===)
        if '=' * 30 in line:
            i += 1
            if i < len(lines):
                heading_text = lines[i].strip()
                if heading_text and '=' not in heading_text:
                    # Main chapter heading
                    if 'CHAPTER' in heading_text.upper() or 'SECTION' in heading_text.upper():
                        doc.add_heading(heading_text, level=0)
                    else:
                        doc.add_heading(heading_text, level=1)
                i += 1
            # Skip closing === line
            while i < len(lines) and '=' * 30 in lines[i]:
                i += 1
            continue
        
        # Section headers like "2.1 ABOUT LANKADEEPA"
        section_match = re.match(r'^(\d+\.[\d.]*)\s+(.+)$', line.strip())
        if section_match and line.strip().isupper() or re.match(r'^(\d+\.[\d.]*)\s+[A-Z][A-Z\s]+$', line.strip()):
            doc.add_heading(line.strip(), level=2)
            i += 1
            continue
        
        # Subsection headers (ending with dashes)
        if line.strip().endswith('---') or line.strip().endswith('---'):
            i += 1
            continue
        
        # Bold subsection headers (line followed by dashes)
        if i + 1 < len(lines) and lines[i + 1].strip().startswith('---'):
            doc.add_heading(line.strip(), level=3)
            i += 2
            continue
        
        # Table-like content (starts with +)
        if line.strip().startswith('+') or line.strip().startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and (lines[i].strip().startswith('+') or lines[i].strip().startswith('|')):
                table_lines.append(lines[i])
                i += 1
            # Add as preformatted text
            table_text = '\n'.join(table_lines)
            p = doc.add_paragraph()
            run = p.add_run(table_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue
        
        # Code-like content (indented lines starting with keywords)
        if line.startswith('    ') or line.startswith('\t'):
            code_lines = [line]
            i += 1
            while i < len(lines) and (lines[i].startswith('    ') or lines[i].startswith('\t') or not lines[i].strip()):
                if lines[i].strip():
                    code_lines.append(lines[i])
                else:
                    break
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue
        
        # Regular paragraph text
        if line.strip():
            # Collect consecutive non-empty lines as one paragraph
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i]
                # Stop if we hit a header or special line
                if '=' * 20 in next_line:
                    break
                if next_line.strip().endswith('---'):
                    break
                if re.match(r'^(\d+\.[\d.]*)\s+[A-Z]', next_line.strip()):
                    break
                if next_line.strip().startswith('+') or next_line.strip().startswith('|'):
                    break
                if not next_line.strip():
                    break
                para_lines.append(next_line)
                i += 1
            
            # Join and add as paragraph
            para_text = '\n'.join(para_lines)
            p = doc.add_paragraph(para_text)
            continue
        
        i += 1
    
    # Save the document
    output_file = "Data_Section_Thesis.docx"
    doc.save(output_file)
    print(f"Successfully converted to: {output_file}")


if __name__ == "__main__":
    convert_data_section_to_docx()
