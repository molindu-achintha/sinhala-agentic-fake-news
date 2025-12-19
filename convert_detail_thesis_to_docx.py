#!/usr/bin/env python3
"""
Convert detailed methodology thesis text to Word document
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_detail_thesis_to_docx():
    """Convert detailed methodology explanation to formatted Word doc."""
    
    with open('methodology_thesis_detailed.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines at start
        if not line.strip():
            i += 1
            continue
            
        # Parse Title Block in the text file
        if 'CHAPTER 4:' in line:
            # Chapter Title
            p = doc.add_heading(line.strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Skip subsequent decorative lines (=) and subtitle lines
            i += 1
            while i < len(lines) and (not lines[i].strip() or '=' in lines[i] or 'Sinhala' in lines[i] or 'December' in lines[i]):
                if lines[i].strip() and '=' not in lines[i]:
                     # Add subtitles (Sinhala Agentic...) as normal centered text
                     p = doc.add_paragraph(lines[i].strip())
                     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
            continue
        
        # Main Section Headers (4.X TITLE)
        # We look for lines followed by "==="
        if i + 1 < len(lines) and lines[i+1].strip().startswith('==='):
            h = doc.add_heading(line.strip(), level=1)
            i += 2 # Skip current line and === line
            continue
            
        # Sub-Section Headers (4.X.X TITLE)
        # We look for lines followed by "---"
        if i + 1 < len(lines) and lines[i+1].strip().startswith('---'):
            h = doc.add_heading(line.strip(), level=2)
            i += 2 # Skip current line and --- line
            continue

        # Regular 4.X Headers if manual underlining wasn't used/consistent
        header_match = re.match(r'^(4\.\d+)\s+(.+)$', line.strip())
        if header_match:
             # Check if it's already handled by the === check
             if i+1 < len(lines) and '=' in lines[i+1]:
                 pass # Will be caught by next loop iteration or logic above? No, logic above might miss if I am here
                 # Actually if I am here, the Previous check failed.
             h = doc.add_heading(line.strip(), level=1)
             i += 1
             # If there is an underline after, skip it
             if i < len(lines) and '=' in lines[i]:
                 i += 1
             continue
             
        # Regular 4.X.X Headers
        subheader_match = re.match(r'^(4\.\d+\.\d+)\s+(.+)$', line.strip())
        if subheader_match:
             h = doc.add_heading(line.strip(), level=2)
             i += 1
             if i < len(lines) and '-' in lines[i]:
                 i += 1
             continue
        
        # Skip standalone separators
        if line.strip().startswith('===') or line.strip().startswith('---'):
            i += 1
            continue

        # Lists
        if line.strip().startswith('* ') or line.strip().startswith('- '):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
            i += 1
            continue
        
        # Numbered lists like "1. Item"
        if re.match(r'^\d+\.\s+', line.strip()):
            text_part = re.sub(r'^\d+\.\s+', '', line.strip())
            p = doc.add_paragraph(text_part, style='List Number')
            i += 1
            continue

        # Code blocks (```)
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue

        # Regular paragraph
        if line.strip():
            # If it looks like a paragraph, accumulate lines
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i]
                if not next_line.strip():
                    break
                if next_line.strip().startswith('='):
                    break
                if next_line.strip().startswith('-'):
                    break
                if next_line.strip().startswith('*'):
                    break
                if re.match(r'^\d+\.\s+', next_line.strip()):
                    break
                if next_line.strip().startswith('```'):
                    break
                if re.match(r'^4\.\d+', next_line.strip()):
                    break
                    
                para_lines.append(next_line.strip())
                i += 1
            
            para_text = ' '.join(para_lines)
            p = doc.add_paragraph(para_text)
            continue
            
        i += 1
    
    output_file = "Methodology_Detailed_Thesis.docx"
    doc.save(output_file)
    print(f"Created: {output_file}")


if __name__ == "__main__":
    convert_detail_thesis_to_docx()
